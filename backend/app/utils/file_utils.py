import os, sys
import csv
from datetime import datetime
import requests
from bs4 import BeautifulSoup
from io import BytesIO
import docx
from docx import Document
import pdfplumber
import numpy as np
import pandas as pd
from sentence_transformers import SentenceTransformer
from PyPDF2 import PdfReader

from app.logger import logger 
from app.exception import CustomException

from dotenv import load_dotenv
load_dotenv()


UPLOAD_DIR = os.getenv("UPLOAD_DIR")
LOG_PATH = os.path.join(UPLOAD_DIR, "file_content_log.csv")

def log_file_data(file_path: str, content: str | pd.DataFrame):
    filename = os.path.basename(file_path)
    timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")

    if isinstance(content, pd.DataFrame):
        content_str = content.head(5).to_string(index=False)  # log preview only
    else:
        content_str = content[:1000]  # Limit large text logs for performance

    # Check if log file exists
    file_exists = os.path.isfile(LOG_PATH)

    with open(LOG_PATH, mode="a", newline="", encoding="utf-8") as f:
        writer = csv.DictWriter(f, fieldnames=["timestamp", "filename", "content_preview"])
        if not file_exists:
            writer.writeheader()
        writer.writerow({
            "timestamp": timestamp,
            "filename": filename,
            "content_preview": content_str
        })

logger.info("Data Extraction of file started")

def read_txt(path):
    try: 
        logger.info(f"Data Extraction of txt {path} file started")
        with open(path, "r", encoding="utf-8") as f:
            log_file_data(path, str(f.read()))
            return f.read()
    except Exception as e: 
        logger.info(f"Exception occure during Extraction of txt {path}") 
        raise CustomException(e, sys)
        
def extract_text_from_pdf(file_path: str) -> str:
    try:
        logger.info(f"Data Extraction of pdf {file_path} file started") 
        text = ""
        reader = PdfReader(file_path)
        for page in reader.pages:
            text += page.extract_text() + "\n"
        log_file_data(file_path, str(text))    
        return text
    except Exception as e: 
        logger.info(f"Exception occure during Extraction of pdf {file_path}") 
        raise CustomException(e, sys)
    
def extract_text_from_docx(file_path: str) -> str:
    try:
        logger.info(f"Data Extraction of doc {file_path} file started") 
        doc = docx.Document(file_path)
        log_file_data(file_path, str("\n".join([para.text for para in doc.paragraphs])))
        return "\n".join([para.text for para in doc.paragraphs])
    except Exception as e: 
        logger.info(f"Exception occure during Extraction of docx {file_path}") 
        raise CustomException(e, sys)    

def extract_text_from_csv(file_path: str):
    try:
        logger.info(f"Data Extraction of csv {file_path} file started")
        df = pd.read_csv(file_path)
        log_file_data(file_path, df) 
        return df
    except Exception as e: 
        logger.info(f"Exception occure during Extraction of csv {file_path}") 
        raise CustomException(e, sys)

def extract_text_from_excel(file_path: str):
    try:
        logger.info(f"Data Extraction of excel {file_path} file started") 
        df = pd.read_excel(file_path)
        log_file_data(file_path, df)
        return df
    except Exception as e: 
        logger.info(f"Exception occure during Extraction of excel {file_path}") 
        raise CustomException(e, sys)
    

def extract_from_url(url: str):
        logger.info(f"URL extraction started: {url}")
        result = {}
        try:
            response = requests.get(url)
            content_type = response.headers.get("Content-Type", "")

            if "text/html" in content_type:
                logger.info("Processing HTML content")
                soup = BeautifulSoup(response.content, "lxml")
                page_text = soup.get_text(separator="\n", strip=True)
                tables = pd.read_html(url)
                result["text"] = page_text
                result["tables"] = [df.to_dict(orient="records") for df in tables]
                # result["tables"] = tables
                logger.info(f"Extracted text from url : {str(page_text[:100])}")
                logger.info(f"No of tables extracted from url : {len(tables)}")

            # elif "application/pdf" in content_type:
            #     logger.info("Processing PDF content")
            #     with pdfplumber.open(BytesIO(response.content)) as pdf:
            #         pdf_text = ""
            #         pdf_tables = []
            #         for page in pdf.pages:
            #             pdf_text += page.extract_text() or ""
            #             tables = page.extract_tables()
            #             for table in tables:
            #                 pdf_tables.append(table)
            #         result["text"] = pdf_text
            #         result["tables"] = pdf_tables

            # elif "application/vnd.openxmlformats-officedocument.wordprocessingml.document" in content_type:
            #     logger.info("Processing DOCX content")
            #     doc = Document(BytesIO(response.content))
            #     doc_text = "\n".join([p.text for p in doc.paragraphs])
            #     tables = []
            #     for table in doc.tables:
            #         table_data = []
            #         for row in table.rows:
            #             table_data.append([cell.text.strip() for cell in row.cells])
            #         tables.append(table_data)
            #     result["text"] = doc_text
            #     result["tables"] = tables

            else:
                logger.warning(f"Unsupported content type: {content_type}")
                raise ValueError(f"Unsupported content type: {content_type}")

        except Exception as e:
            logger.error(f"Error fetching or processing URL: {e}")
            raise ValueError(f"Error fetching or processing URL: {e}")

        return result





